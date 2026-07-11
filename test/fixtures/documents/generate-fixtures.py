#!/usr/bin/env python3
"""Generate synthetic test documents for Mike Atlas local QA."""
import os
from pathlib import Path

ROOT = Path(__file__).parent
WATERMARK = "SYNTHETIC TEST DOCUMENT — NO REAL CLIENT DATA"


def pdf_bytes(text_lines, pages=1):
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
    except ImportError:
        print("reportlab not installed; skipping PDF generation")
        return None

    from io import BytesIO
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    width, height = letter
    for _ in range(pages):
        y = height - 72
        c.setFont("Helvetica", 12)
        for line in text_lines:
            c.drawString(72, y, line)
            y -= 18
        c.setFont("Helvetica-Bold", 14)
        c.drawString(72, height - 144, WATERMARK)
        c.showPage()
    c.save()
    return buf.getvalue()


def make_docx(text):
    try:
        from docx import Document
    except ImportError:
        print("python-docx not installed; skipping DOCX generation")
        return None
    doc = Document()
    doc.add_heading("Synthetic Contract", level=1)
    doc.add_paragraph(text)
    doc.add_paragraph(WATERMARK)
    from io import BytesIO
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_xlsx():
    try:
        from openpyxl import Workbook
    except ImportError:
        print("openpyxl not installed; skipping XLSX generation")
        return None
    wb = Workbook()
    ws = wb.active
    ws.title = "Synthetic"
    ws["A1"] = "Item"
    ws["B1"] = "Value"
    ws["A2"] = "Test"
    ws["B2"] = 123
    ws["A3"] = WATERMARK
    from io import BytesIO
    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


def main():
    (ROOT / "sample-contract.pdf").write_bytes(
        pdf_bytes(["This is a synthetic contract for testing.", "Paragraph two."]) or b""
    )
    docx = make_docx("This is a synthetic DOCX contract for local testing.")
    if docx:
        (ROOT / "sample-contract.docx").write_bytes(docx)
    xlsx = make_xlsx()
    if xlsx:
        (ROOT / "sample-spreadsheet.xlsx").write_bytes(xlsx)
    (ROOT / "sample-nda.pdf").write_bytes(
        pdf_bytes(["NON-DISCLOSURE AGREEMENT (SYNTHETIC)", "Parties: Alpha and Beta"]) or b""
    )
    (ROOT / "empty.pdf").write_bytes(pdf_bytes([]) or b"")
    (ROOT / "invalid-extension.txt").write_text(
        f"{WATERMARK}\nThis file has an invalid extension for upload tests."
    )
    # Corrupted PDF: take a valid PDF and truncate it
    full = pdf_bytes(["This PDF will be corrupted for error handling tests."])
    if full:
        (ROOT / "corrupted.pdf").write_bytes(full[:64])
    # Near-limit file (e.g., 4 MB)
    near = pdf_bytes(["Large synthetic document."] * 1000, pages=50)
    if near:
        (ROOT / "near-limit.pdf").write_bytes(near)
    print("Fixtures generated in", ROOT)


if __name__ == "__main__":
    main()
